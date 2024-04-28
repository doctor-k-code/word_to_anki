"""
[name] app.py
[purpose] word to csv for anki
[reference]
    https://docs.streamlit.io/develop/api-reference

written by Dr.K, 2024/4/28
"""
import datetime
import glob
from io import StringIO
import os
import shutil

import docx
import numpy as np
import pandas as pd
import streamlit as st


st.title("wordファイルをanki用のCSVファイルに変換します")

uploaded_files = st.file_uploader(
    "wordファイルをアップロード", type="docx", accept_multiple_files=True)

table = None
for uploaded_file in uploaded_files:
    st.write("filename:", uploaded_file.name)
    # bytes_data = uploaded_file.getvalue()
    # st.write(bytes_data)
    # stringio = StringIO(bytes_data.decode("utf-8"))
    # st.write(stringio)
    doc = docx.Document(uploaded_file)
    st.write(doc.paragraphs)
"""
    # 空白を除き、各段落をリストで管理する。
    all_paragraph = []
    for par in doc.paragraphs:
        text = par.text
        if text:
            all_paragraph.append(text)

    # 表（front）、裏（back）、ページ数（page）のペアになるはずである
    # 3の倍数になっていることを確認する
    rest_par_num = len(all_paragraph) % 3
    if rest_par_num != 0:
        # どこかに過不足があるためにズレが生じていると考えられる
        # 余りを除いて表を作成し、Debugの参考にする
        all_paragraph = all_paragraph[:-rest_par_num]
        st.write(f'段落数が一致しません。表を確認してください: {uploaded_file.name}')

    st.write(all_paragraph[0])"""
